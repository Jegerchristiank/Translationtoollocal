from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any


NUMBER_COL_TWIPS = 601
GAP_COL_TWIPS = 329
PAGE_WIDTH_TWIPS = 11906
SIDE_MARGIN_TWIPS = 1134
TEXT_COL_TWIPS = PAGE_WIDTH_TWIPS - (SIDE_MARGIN_TWIPS * 2) - NUMBER_COL_TWIPS - GAP_COL_TWIPS


def _header_date(job: dict[str, Any]) -> str:
    created_at = str(job.get("created_at") or "").strip()
    if created_at:
        try:
            parsed = datetime.fromisoformat(created_at.replace("Z", "+00:00"))
            if parsed.tzinfo is not None:
                parsed = parsed.astimezone()
            return parsed.strftime("%d.%m.%Y")
        except ValueError:
            pass
    return datetime.now().strftime("%d.%m.%Y")


def _source_label(job: dict[str, Any]) -> str:
    source_name = str(job.get("source_name") or "").strip()
    if source_name:
        return Path(source_name).stem
    return Path(str(job.get("source_path", ""))).stem


def _header_lines(job: dict[str, Any]) -> list[str]:
    source_name = _source_label(job)
    duration_sec = float(job.get("duration_sec", 0) or 0)
    duration_min = max(1, round(duration_sec / 60))
    date_str = _header_date(job)

    return [
        f'Navn på fil: "{source_name}"',
        f"Dato: {date_str}",
        f"Varighed: {duration_min} minutter",
        "",
        "Deltagere:",
        "Interviewer (I)",
        "Deltager (D)",
        "",
    ]


def _line_entries(segments: list[dict[str, Any]]) -> list[tuple[int, str, str]]:
    entries: list[tuple[int, str, str]] = []
    line_no = 1
    for segment in segments:
        speaker = str(segment.get("speaker") or "D")
        text = str(segment.get("text") or "").strip()
        if not text:
            continue
        entries.append((line_no, speaker, text))
        line_no += 1
    return entries


def export_txt(job: dict[str, Any], transcript: list[dict[str, Any]], output_path: Path) -> None:
    lines = _header_lines(job)

    for line_no, speaker, text in _line_entries(transcript):
        lines.append(f"{line_no}\t{speaker}: {text}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def export_docx(job: dict[str, Any], transcript: list[dict[str, Any]], output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Mm, Pt, Twips
    except ImportError as exc:  # pragma: no cover - env dependent
        raise RuntimeError("python-docx mangler. Installer python/requirements.txt") from exc

    def _format_paragraph(paragraph: Any) -> None:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

    def _set_white_borders(cell: Any) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)

        for edge in ("top", "left", "bottom", "right"):
            edge_tag = qn(f"w:{edge}")
            edge_element = tc_borders.find(edge_tag)
            if edge_element is None:
                edge_element = OxmlElement(f"w:{edge}")
                tc_borders.append(edge_element)
            edge_element.set(qn("w:val"), "single")
            edge_element.set(qn("w:sz"), "4")
            edge_element.set(qn("w:space"), "0")
            edge_element.set(qn("w:color"), "FFFFFF")

    def _append_table_row(table: Any, line_no: int, text: str, speaker: str | None = None) -> int:
        row = table.add_row()
        row.height = Cm(0.5)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.cells[0].width = Twips(NUMBER_COL_TWIPS)
        row.cells[1].width = Twips(GAP_COL_TWIPS)
        row.cells[2].width = Twips(TEXT_COL_TWIPS)

        for cell in row.cells:
            _set_white_borders(cell)

        number_p = row.cells[0].paragraphs[0]
        _format_paragraph(number_p)
        number_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        number_p.add_run(str(line_no))

        gap_p = row.cells[1].paragraphs[0]
        _format_paragraph(gap_p)

        text_p = row.cells[2].paragraphs[0]
        _format_paragraph(text_p)
        if speaker:
            speaker_run = text_p.add_run(f"{speaker}:")
            speaker_run.bold = True
            if text:
                text_p.add_run(f" {text}")
        elif text:
            text_p.add_run(text)

        return line_no + 1

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(30)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    headers = _header_lines(job)
    for line in headers:
        paragraph = doc.add_paragraph(line)
        _format_paragraph(paragraph)
        if line == "Deltagere:" and paragraph.runs:
            paragraph.runs[0].bold = True

    line_entries = _line_entries(transcript)
    if line_entries:
        table = doc.add_table(rows=0, cols=3)
        table.style = "Table Grid"
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Twips(NUMBER_COL_TWIPS)
        table.columns[1].width = Twips(GAP_COL_TWIPS)
        table.columns[2].width = Twips(TEXT_COL_TWIPS)

        for number, speaker, text in line_entries:
            _append_table_row(table, number, text, speaker=speaker)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
