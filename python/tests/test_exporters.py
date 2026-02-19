from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm, Mm, Pt, Twips

from transkriptor.exporters import TEXT_COL_TWIPS, export_docx, export_txt


def test_export_txt_contains_header_and_dialogue(tmp_path: Path):
    job = {
        "source_path": "/tmp/demo_interview.wav",
        "duration_sec": 3600,
    }
    transcript = [
        {"speaker": "I", "text": "Velkommen", "startSec": 0, "endSec": 1},
        {"speaker": "D", "text": "Tak", "startSec": 1.1, "endSec": 1.6},
    ]

    out = tmp_path / "out.txt"
    export_txt(job, transcript, out)
    content = out.read_text(encoding="utf-8")

    assert "Navn på fil:" in content
    assert "Varighed:" in content
    assert "1\tI: Velkommen" in content
    assert "2\tD: Tak" in content


def test_export_docx_uses_numbered_table_layout(tmp_path: Path):
    job = {
        "source_path": "/tmp/demo_interview.wav",
        "duration_sec": 1200,
        "created_at": "2026-02-11T10:15:00+00:00",
    }
    transcript = [
        {
            "speaker": "I",
            "text": " ".join(["ord"] * 120),
            "startSec": 0,
            "endSec": 25,
        },
        {
            "speaker": "D",
            "text": "Det giver mening.",
            "startSec": 30,
            "endSec": 33,
        },
    ]

    out = tmp_path / "out.docx"
    export_docx(job, transcript, out)

    doc = Document(out)
    assert doc.paragraphs[0].text.startswith('Navn på fil: "demo_interview"')
    assert any(p.text == "Deltagere:" for p in doc.paragraphs)
    assert len(doc.tables) == 1

    table = doc.tables[0]
    assert len(table.columns) == 3
    assert len(table.rows) == 2
    assert table.rows[0].cells[0].text == "1"
    assert table.rows[0].cells[2].text.startswith("I:")
    assert table.rows[1].cells[0].text == "2"
    assert table.rows[1].cells[2].text.startswith("D:")

    numbers = [int(row.cells[0].text) for row in table.rows]
    assert numbers == list(range(1, len(table.rows) + 1))
    assert abs(int(table.columns[2].width) - int(Twips(TEXT_COL_TWIPS))) <= 80

    first_row = table.rows[0]
    assert first_row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY
    assert abs(int(first_row.height) - int(Cm(0.5))) <= 400

    section = doc.sections[0]
    assert abs(int(section.left_margin) - int(Mm(20))) <= 400
    assert abs(int(section.right_margin) - int(Mm(20))) <= 400

    normal_style = doc.styles["Normal"]
    assert abs(int(normal_style.font.size) - int(Pt(12))) <= 20
